from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
from datetime import datetime
import hashlib
import hmac
import json
import io
import os

OUTPUT_DIR = "generated_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def generate_key_facts_token(policy_number: str) -> str:
    """Short HMAC token that authorizes a guest (unauthenticated) download of
    one policy's Key Facts document. Derived from the policy number and the
    JWT secret, so it cannot be guessed from the policy number alone."""
    from app.core.config import settings

    return hmac.new(
        settings.jwt_secret_key.encode(),
        policy_number.encode(),
        hashlib.sha256,
    ).hexdigest()[:16]


def verify_key_facts_token(policy_number: str, token: str) -> bool:
    import hmac as _hmac

    return _hmac.compare_digest(token, generate_key_facts_token(policy_number))


def generate_key_facts_pdf(
    product_type: str,
    coverage_blocks: list[str],
    premium_monthly: float,
    premium_annual: float,
    policy_number: str,
    holder_name: str,
    sla_data: list[dict] = None,
) -> str:
    """Generate a 'Key Facts' PDF document for a policy."""
    pdf = FPDF()
    pdf.add_page()

    # Header
    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 15, "InsurBridge AI", ln=True, align="C")
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Key Facts Document", ln=True, align="C")
    pdf.ln(5)

    # Line separator
    pdf.set_draw_color(0, 102, 204)
    pdf.set_line_width(0.5)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    # Policy Info
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Policy Details", ln=True)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, f"Policy Number: {policy_number}", ln=True)
    pdf.cell(0, 6, f"Policyholder: {holder_name}", ln=True)
    pdf.cell(0, 6, f"Product: {product_type}", ln=True)
    pdf.cell(0, 6, f"Date: {datetime.now().strftime('%d %B %Y')}", ln=True)
    pdf.ln(5)

    # Coverage
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Your Selected Coverage", ln=True)
    pdf.set_font("Arial", "", 10)
    for block in coverage_blocks:
        pdf.cell(0, 6, f"  - {block}", ln=True)
    pdf.ln(5)

    # Premium
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Premium Summary", ln=True)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, f"Monthly Premium: {premium_monthly:.2f}", ln=True)
    pdf.cell(0, 6, f"Annual Premium: {premium_annual:.2f}", ln=True)
    pdf.ln(5)

    # SLAs
    if sla_data:
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Service Level Commitments", ln=True)
        pdf.set_font("Arial", "", 10)
        for sla in sla_data:
            pdf.cell(0, 6, f"  - {sla['metric']}: {sla['promised']}", ln=True)

    # Save
    filename = f"{OUTPUT_DIR}/key_facts_{policy_number}.pdf"
    pdf.output(filename)
    return filename


def generate_key_facts_docx(
    product_type: str,
    coverage_blocks: list[str],
    premium_monthly: float,
    premium_annual: float,
    policy_number: str,
    holder_name: str,
    sla_data: list[dict] = None,
) -> str:
    """Generate a Key Facts DOCX document for a policy."""
    doc = Document()

    # Title
    title = doc.add_heading("InsurBridge AI", level=0)
    doc.add_heading("Key Facts Document", level=1)

    # Policy Details
    doc.add_heading("Policy Details", level=2)
    doc.add_paragraph(f"Policy Number: {policy_number}")
    doc.add_paragraph(f"Policyholder: {holder_name}")
    doc.add_paragraph(f"Product: {product_type}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")

    # Coverage
    doc.add_heading("Your Selected Coverage", level=2)
    for block in coverage_blocks:
        doc.add_paragraph(block, style="List Bullet")

    # Premium
    doc.add_heading("Premium Summary", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Monthly"
    table.cell(0, 1).text = f"{premium_monthly:.2f}"
    table.cell(1, 0).text = "Annual"
    table.cell(1, 1).text = f"{premium_annual:.2f}"

    # SLAs
    if sla_data:
        doc.add_heading("Service Level Commitments", level=2)
        for sla in sla_data:
            doc.add_paragraph(f"{sla['metric']}: {sla['promised']}", style="List Bullet")

    # Save
    filename = f"{OUTPUT_DIR}/key_facts_{policy_number}.docx"
    doc.save(filename)
    return filename


def generate_sla_report_pdf(
    tenant_id: str,
    sla_records: list[dict],
    report_period: str = "Monthly",
) -> str:
    """Generate an SLA Performance Report as PDF."""
    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Arial", "B", 18)
    pdf.cell(0, 12, "SLA Performance Report", ln=True, align="C")
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, f"Tenant: {tenant_id}", ln=True, align="C")
    pdf.cell(0, 6, f"Period: {report_period}", ln=True, align="C")
    pdf.cell(0, 6, f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}", ln=True, align="C")
    pdf.ln(10)

    # Table Header
    pdf.set_font("Arial", "B", 10)
    pdf.cell(60, 8, "Metric", border=1)
    pdf.cell(40, 8, "Promised", border=1)
    pdf.cell(40, 8, "Actual", border=1)
    pdf.cell(30, 8, "Status", border=1)
    pdf.ln()

    # Table Rows
    pdf.set_font("Arial", "", 9)
    for record in sla_records:
        pdf.cell(60, 7, str(record.get("metric", "")), border=1)
        pdf.cell(40, 7, str(record.get("promised", "")), border=1)
        pdf.cell(40, 7, str(record.get("actual", "N/A")), border=1)
        status = "BREACH" if record.get("breached") else "OK"
        pdf.cell(30, 7, status, border=1)
        pdf.ln()

    filename = f"{OUTPUT_DIR}/sla_report_{tenant_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    pdf.output(filename)
    return filename
